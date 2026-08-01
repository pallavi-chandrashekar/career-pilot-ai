"""Render approved application package content to DOCX and PDF bytes."""

from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER  # type: ignore[import-untyped]
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-untyped]


def render_docx(*, title: str, content: dict[str, object]) -> bytes:
    document = Document()
    document.add_heading(title, level=0)
    _add_docx_section(document, "Selected experience", content.get("tailored_resume", []))
    _add_docx_section(document, "Cover letter", content.get("cover_letter", ""))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(*, title: str, content: dict[str, object]) -> bytes:
    output = BytesIO()
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for heading, value in (
        ("Selected experience", content.get("tailored_resume", [])),
        ("Cover letter", content.get("cover_letter", "")),
    ):
        story.append(Paragraph(heading, styles["Heading2"]))
        for paragraph in _paragraphs(value):
            story.append(Paragraph(paragraph.replace("&", "&amp;"), styles["BodyText"]))
            story.append(Spacer(1, 6))
    SimpleDocTemplate(output, pagesize=LETTER).build(story)
    return output.getvalue()


def _paragraphs(value: object) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value]
    return [line for line in str(value).splitlines() if line]


def _add_docx_section(document: DocxDocument, heading: str, value: object) -> None:
    document.add_heading(heading, level=1)
    for paragraph in _paragraphs(value):
        document.add_paragraph(paragraph)
