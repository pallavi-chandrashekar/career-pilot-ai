"""Real HTTP document upload coverage using fictional data and the Compose stack."""

import os
from io import BytesIO
from uuid import uuid4

import httpx
import pytest
from docx import Document as DocxDocument

API_BASE_URL = os.getenv("E2E_API_BASE_URL")
pytestmark = pytest.mark.e2e


def _fictional_docx() -> bytes:
    document = DocxDocument()
    document.add_paragraph("SUMMARY")
    document.add_paragraph("Fictional Candidate")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.skipif(not API_BASE_URL, reason="E2E_API_BASE_URL is required for document E2E tests")
def test_document_upload_deduplication_and_owner_isolation() -> None:
    email = f"fictional-document-{uuid4().hex}@example.com"
    registration = httpx.post(
        f"{API_BASE_URL}/api/v1/auth/register",
        json={
            "email": email,
            "display_name": "Fictional Candidate",
            "timezone": "UTC",
            "password": "fictional-password-123",
        },
        timeout=10.0,
    )
    assert registration.status_code == 201
    headers = {"Authorization": f"Bearer {registration.json()['access_token']}"}
    files = {
        "file": (
            "fictional-resume.docx",
            _fictional_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    upload = httpx.post(
        f"{API_BASE_URL}/api/v1/documents", headers=headers, files=files, timeout=10.0
    )
    assert upload.status_code == 201
    document = upload.json()
    assert document["status"] == "UPLOADED"
    assert "storage_key" not in document

    duplicate = httpx.post(
        f"{API_BASE_URL}/api/v1/documents", headers=headers, files=files, timeout=10.0
    )
    assert duplicate.status_code == 200
    assert duplicate.json()["id"] == document["id"]

    parsed = httpx.post(
        f"{API_BASE_URL}/api/v1/documents/{document['id']}/parse", headers=headers, timeout=10.0
    )
    assert parsed.status_code == 200
    assert parsed.json()["status"] == "PARSED"
    assert parsed.json()["parser_version"] == "1"

    extraction = httpx.post(
        f"{API_BASE_URL}/api/v1/documents/{document['id']}/extract-claims",
        headers=headers,
        json={"provider": "openai"},
        timeout=60.0,
    )
    assert extraction.status_code == 200
    assert extraction.json()["status"] == "COMPLETED"
    repeated_extraction = httpx.post(
        f"{API_BASE_URL}/api/v1/documents/{document['id']}/extract-claims",
        headers=headers,
        json={"provider": "openai"},
        timeout=10.0,
    )
    assert repeated_extraction.status_code == 200
    assert repeated_extraction.json()["id"] == extraction.json()["id"]
    claims = httpx.get(f"{API_BASE_URL}/api/v1/candidate-claims", headers=headers, timeout=10.0)
    assert claims.status_code == 200
    assert all(claim["verification_status"] == "DRAFT" for claim in claims.json())

    second_registration = httpx.post(
        f"{API_BASE_URL}/api/v1/auth/register",
        json={
            "email": f"fictional-owner-{uuid4().hex}@example.com",
            "display_name": "Second Fictional Candidate",
            "timezone": "UTC",
            "password": "fictional-password-123",
        },
        timeout=10.0,
    )
    second_headers = {"Authorization": f"Bearer {second_registration.json()['access_token']}"}
    second_claims = httpx.get(
        f"{API_BASE_URL}/api/v1/candidate-claims", headers=second_headers, timeout=10.0
    )
    assert second_claims.status_code == 200
    assert second_claims.json() == []
    hidden = httpx.get(
        f"{API_BASE_URL}/api/v1/documents/{document['id']}/status",
        headers=second_headers,
        timeout=10.0,
    )
    assert hidden.status_code == 404
